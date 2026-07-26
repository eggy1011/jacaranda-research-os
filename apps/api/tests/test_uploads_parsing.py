from __future__ import annotations

from collections.abc import AsyncIterator
from pathlib import Path
from typing import Any

import httpx
import pytest
from auth_helpers import sign_in
from pytest import MonkeyPatch
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from jacaranda_api.config import get_settings
from jacaranda_api.db.engine import create_engine, create_session_factory
from jacaranda_api.db.models import Base, Upload
from jacaranda_api.documents.parser import (
    DocumentParseError,
    UnsupportedDocumentError,
    parse_document,
)
from jacaranda_api.main import create_app
from jacaranda_api.pipeline.upload_evidence import merge_upload_evidence
from jacaranda_api.worker import parse_upload


def _write_pdf(path: Path) -> None:
    """Build a minimal one-page PDF (Helvetica text + valid xref) by hand."""
    content = b"BT /F1 12 Tf 72 720 Td (Revenue FY2025: 5000 million CNY) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R"
        b" /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >> stream\n%s\nendstream" % (len(content), content),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    body = b"%PDF-1.4\n"
    offsets = []
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += b"%d 0 obj\n%s\nendobj\n" % (number, obj)
    xref_offset = len(body)
    body += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        body += b"%010d 00000 n \n" % offset
    body += (
        b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n"
        % (len(objects) + 1, xref_offset)
    )
    path.write_bytes(body)


def _write_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("2025年度经营回顾")
    document.add_paragraph("公司实现营业收入50亿元。")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "数值"
    table.rows[1].cells[0].text = "毛利率"
    table.rows[1].cells[1].text = "45%"
    document.save(str(path))


def _write_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "财务"
    sheet.append(["项目", "2025"])
    sheet.append(["营业收入", 5_000_000_000])
    workbook.save(str(path))


class TestParser:
    def test_pdf_pages_with_locators(self, tmp_path: Path) -> None:
        path = tmp_path / "report.pdf"
        _write_pdf(path)
        parsed = parse_document(path)
        assert len(parsed.blocks) == 1
        assert parsed.blocks[0].locator == "page=1"
        assert "Revenue FY2025" in parsed.blocks[0].text

    def test_docx_paragraphs_and_tables(self, tmp_path: Path) -> None:
        path = tmp_path / "review.docx"
        _write_docx(path)
        parsed = parse_document(path)
        locators = [block.locator for block in parsed.blocks]
        assert "paragraph=1" in locators
        assert "table=1" in locators
        table_block = next(block for block in parsed.blocks if block.kind == "table")
        assert "毛利率\t45%" in table_block.text

    def test_xlsx_sheets(self, tmp_path: Path) -> None:
        path = tmp_path / "data.xlsx"
        _write_xlsx(path)
        parsed = parse_document(path)
        assert parsed.blocks[0].locator == "sheet=财务"
        assert "营业收入\t5000000000" in parsed.blocks[0].text

    def test_unknown_suffix_rejected(self, tmp_path: Path) -> None:
        path = tmp_path / "notes.txt"
        path.write_text("hello", encoding="utf-8")
        with pytest.raises(UnsupportedDocumentError):
            parse_document(path)

    def test_corrupt_pdf_fails_safely(self, tmp_path: Path) -> None:
        path = tmp_path / "broken.pdf"
        path.write_bytes(b"%PDF-1.4 not really")
        with pytest.raises(DocumentParseError):
            parse_document(path)


class TestUploadEvidenceMerge:
    def _evidence(self) -> dict[str, Any]:
        return {
            "sources": [
                {
                    "source_id": "SRC-001",
                    "type": "market_data_api",
                    "title": "quote",
                    "url_or_document": "provider://akshare/quote/X",
                    "retrieved_at": "2026-07-27T09:00:00Z",
                    "reliability_tier": "primary",
                }
            ],
            "warnings": [],
        }

    def test_sources_continue_numbering_and_chunks_carry_locators(self) -> None:
        upload = {
            "upload_id": "abc123",
            "filename": "annual.pdf",
            "created_at": "2026-07-27T08:00:00+00:00",
            "blocks": [
                {"locator": "page=1", "kind": "text", "text": "第一页内容"},
                {"locator": "page=2", "kind": "text", "text": "第二页内容"},
            ],
        }
        evidence, chunks = merge_upload_evidence(self._evidence(), [upload])
        assert [source["source_id"] for source in evidence["sources"]] == ["SRC-001", "SRC-002"]
        upload_source = evidence["sources"][1]
        assert upload_source["type"] == "user_upload"
        assert upload_source["reliability_tier"] == "secondary"
        assert upload_source["url_or_document"] == "upload://abc123"
        assert [chunk["url_or_document"] for chunk in chunks] == [
            "upload://abc123#page=1",
            "upload://abc123#page=2",
        ]

    def test_budget_truncation_produces_warning(self) -> None:
        upload = {
            "upload_id": "big",
            "filename": "big.pdf",
            "created_at": "2026-07-27T08:00:00+00:00",
            "blocks": [
                {"locator": f"page={index}", "kind": "text", "text": "x" * 1000}
                for index in range(1, 30)
            ],
        }
        evidence, chunks = merge_upload_evidence(self._evidence(), [upload])
        assert len(chunks) <= 12
        assert any("truncated" in warning for warning in evidence["warnings"])

    def test_no_uploads_is_identity(self) -> None:
        evidence = self._evidence()
        merged, chunks = merge_upload_evidence(evidence, [])
        assert merged is evidence
        assert chunks == []


@pytest.fixture
async def db(tmp_path: Path) -> AsyncIterator[async_sessionmaker[AsyncSession]]:
    engine = create_engine(f"sqlite:///{tmp_path}/test.db")
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    yield create_session_factory(engine)
    await engine.dispose()


class FakeQueue:
    def __init__(self) -> None:
        self.jobs: list[tuple[str, tuple[object, ...]]] = []

    async def enqueue(self, function: str, *args: object) -> None:
        self.jobs.append((function, args))


@pytest.fixture
async def client(
    db: async_sessionmaker[AsyncSession],
    tmp_path: Path,
    monkeypatch: MonkeyPatch,
) -> AsyncIterator[tuple[httpx.AsyncClient, FakeQueue]]:
    monkeypatch.setenv("DATABASE_URL", "sqlite:///ignored.db")
    monkeypatch.setenv("REDIS_URL", "redis://ignored")
    monkeypatch.setenv("DATA_DIR", str(tmp_path / "data"))
    get_settings.cache_clear()
    app = create_app("test")
    queue = FakeQueue()
    app.state.session_factory = db
    app.state.job_queue = queue
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as http:
        await sign_in(http, db)
        yield http, queue
    get_settings.cache_clear()


@pytest.mark.anyio
async def test_upload_flow_stores_parses_and_reports(
    client: tuple[httpx.AsyncClient, FakeQueue],
    db: async_sessionmaker[AsyncSession],
    tmp_path: Path,
) -> None:
    http, queue = client
    project_id = (await http.post("/projects", json={"symbol": "600519"})).json()["id"]

    docx_path = tmp_path / "review.docx"
    _write_docx(docx_path)
    response = await http.post(
        f"/projects/{project_id}/uploads",
        files={
            "file": (
                "review.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202
    upload_id = response.json()["id"]
    assert queue.jobs == [("parse_upload", (upload_id,))]

    # worker parses it
    outcome = await parse_upload({"session_factory": db}, upload_id)
    assert outcome == "parsed"
    detail = (await http.get(f"/uploads/{upload_id}")).json()
    assert detail["status"] == "parsed"
    async with db() as session:
        upload = await session.get(Upload, upload_id)
        assert upload is not None
        assert upload.parsed is not None
        locators = [block["locator"] for block in upload.parsed["blocks"]]
        assert "paragraph=1" in locators

    listed = (await http.get(f"/projects/{project_id}/uploads")).json()
    assert [item["id"] for item in listed] == [upload_id]


@pytest.mark.anyio
async def test_upload_rejects_bad_type_and_empty(
    client: tuple[httpx.AsyncClient, FakeQueue],
) -> None:
    http, _ = client
    project_id = (await http.post("/projects", json={"symbol": "600519"})).json()["id"]
    bad = await http.post(
        f"/projects/{project_id}/uploads",
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )
    assert bad.status_code == 422
    empty = await http.post(
        f"/projects/{project_id}/uploads",
        files={"file": ("empty.pdf", b"", "application/pdf")},
    )
    assert empty.status_code == 422
    missing = await http.post(
        "/projects/nope/uploads",
        files={"file": ("a.pdf", b"x", "application/pdf")},
    )
    assert missing.status_code == 404


@pytest.mark.anyio
async def test_parse_upload_marks_corrupt_file_failed(
    client: tuple[httpx.AsyncClient, FakeQueue],
    db: async_sessionmaker[AsyncSession],
) -> None:
    http, _ = client
    project_id = (await http.post("/projects", json={"symbol": "600519"})).json()["id"]
    response = await http.post(
        f"/projects/{project_id}/uploads",
        files={"file": ("broken.pdf", b"%PDF-1.4 nope", "application/pdf")},
    )
    upload_id = response.json()["id"]
    outcome = await parse_upload({"session_factory": db}, upload_id)
    assert outcome == "failed"
    detail = (await http.get(f"/uploads/{upload_id}")).json()
    assert detail["status"] == "failed"
    assert detail["error"]["code"] == "DocumentParseError"
