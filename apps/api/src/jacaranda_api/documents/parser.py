from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

ALLOWED_SUFFIXES = {".pdf", ".docx", ".xlsx"}
_MAX_BLOCK_CHARS = 4000


class UnsupportedDocumentError(ValueError):
    def __init__(self, suffix: str) -> None:
        self.suffix = suffix
        super().__init__(f"unsupported document type: {suffix or '(none)'}")


class DocumentParseError(ValueError):
    """The file could not be parsed; the message is safe to store and display."""


@dataclass(frozen=True, slots=True)
class DocumentBlock:
    """One addressable piece of an uploaded document.

    The locator is the fragment appended to ``upload://{file_id}`` so every
    number extracted from an upload can be traced back to a page, paragraph
    or sheet — the provenance rule that applies to provider data applies to
    uploads too.
    """

    locator: str
    kind: str  # "text" | "table"
    text: str


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    blocks: tuple[DocumentBlock, ...]
    warnings: tuple[str, ...] = field(default=())

    def as_dict(self) -> dict[str, object]:
        return {
            "blocks": [
                {"locator": block.locator, "kind": block.kind, "text": block.text}
                for block in self.blocks
            ],
            "warnings": list(self.warnings),
        }


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".xlsx":
        return _parse_xlsx(path)
    raise UnsupportedDocumentError(suffix)


def _clip(text: str) -> str:
    return text if len(text) <= _MAX_BLOCK_CHARS else text[:_MAX_BLOCK_CHARS]


def _parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(path)
        blocks: list[DocumentBlock] = []
        warnings: list[str] = []
        for number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(
                    DocumentBlock(locator=f"page={number}", kind="text", text=_clip(text))
                )
            else:
                warnings.append(f"page {number} contains no extractable text (scanned image?)")
        return ParsedDocument(blocks=tuple(blocks), warnings=tuple(warnings))
    except PyPdfError as error:
        raise DocumentParseError(f"PDF could not be read: {type(error).__name__}") from None


def _parse_docx(path: Path) -> ParsedDocument:
    import docx
    from docx.opc.exceptions import OpcError

    try:
        document = docx.Document(str(path))
    except (OpcError, KeyError, ValueError) as error:
        raise DocumentParseError(f"Word file could not be read: {type(error).__name__}") from None
    blocks: list[DocumentBlock] = []
    for number, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            blocks.append(
                DocumentBlock(locator=f"paragraph={number}", kind="text", text=_clip(text))
            )
    for number, table in enumerate(document.tables, start=1):
        rows = [
            "\t".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
        ]
        text = "\n".join(row for row in rows if row.strip())
        if text:
            blocks.append(DocumentBlock(locator=f"table={number}", kind="table", text=_clip(text)))
    return ParsedDocument(blocks=tuple(blocks))


def _parse_xlsx(path: Path) -> ParsedDocument:
    from openpyxl import load_workbook
    from openpyxl.utils.exceptions import InvalidFileException

    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except (InvalidFileException, KeyError, ValueError) as error:
        raise DocumentParseError(f"Excel file could not be read: {type(error).__name__}") from None
    blocks: list[DocumentBlock] = []
    warnings: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value) for value in row]
                if any(cell.strip() for cell in cells):
                    rows.append("\t".join(cells))
                if len(rows) >= 200:
                    warnings.append(f"sheet {sheet.title} truncated at 200 rows")
                    break
            if rows:
                blocks.append(
                    DocumentBlock(
                        locator=f"sheet={sheet.title}",
                        kind="table",
                        text=_clip("\n".join(rows)),
                    )
                )
    finally:
        workbook.close()
    return ParsedDocument(blocks=tuple(blocks), warnings=tuple(warnings))
